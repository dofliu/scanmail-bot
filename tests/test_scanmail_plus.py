"""ScanMail+ 全面測試 — 新功能覆蓋"""
import io
import os
import sys
import json
import pytest
import numpy as np

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


# ══════════════════════════════════════════════════════════════
# 測試素材工廠
# ══════════════════════════════════════════════════════════════

def _jpg(w=400, h=300, color=(200, 220, 240)):
    import cv2
    img = np.full((h, w, 3), color, dtype=np.uint8)
    cv2.putText(img, "Test", (50, h//2), cv2.FONT_HERSHEY_SIMPLEX, 1, (0,0,0), 2)
    _, buf = cv2.imencode('.jpg', img)
    return buf.tobytes()

def _pdf(text="Test PDF"):
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.drawString(100, 700, text)
    c.save()
    return buf.getvalue()

def _docx(text="Test Document"):
    from docx import Document
    doc = Document()
    doc.add_heading(text, level=1)
    doc.add_paragraph("Test paragraph content.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════
# 文件掃描偵測
# ══════════════════════════════════════════════════════════════

class TestDocScanner:

    def test_detect_document_with_clear_boundary(self):
        import cv2
        from app.services.doc_scanner import detect_document_edges
        img = np.random.randint(60, 130, (500, 400, 3), dtype=np.uint8)
        cv2.rectangle(img, (60, 50), (340, 440), (240, 240, 235), -1)
        _, buf = cv2.imencode('.jpg', img)
        corners = detect_document_edges(buf.tobytes())
        assert corners is not None
        assert len(corners) == 4

    def test_detect_rejects_full_image(self):
        """幾乎整張圖的白色不應被偵測為文件"""
        import cv2
        from app.services.doc_scanner import _is_valid_doc_quad
        # 直接測試驗證邏輯：幾乎整張圖的四邊形應被拒絕
        full_corners = np.array([[5, 5], [795, 5], [795, 595], [5, 595]])
        assert not _is_valid_doc_quad(full_corners, 800, 600)

    def test_perspective_transform(self):
        import cv2
        from app.services.doc_scanner import perspective_transform
        img = np.full((600, 800, 3), 200, dtype=np.uint8)
        _, buf = cv2.imencode('.jpg', img)
        corners = [[100, 50], [700, 80], [680, 550], [120, 520]]
        result = perspective_transform(buf.tobytes(), corners)
        assert len(result) > 0
        decoded = cv2.imdecode(np.frombuffer(result, np.uint8), cv2.IMREAD_COLOR)
        assert decoded is not None

    def test_apply_all_filters(self):
        from app.services.doc_scanner import apply_filter
        data = _jpg(400, 300)
        for f in ["auto", "scan", "color_doc", "document", "enhance", "bw", "original"]:
            result = apply_filter(data, f)
            assert len(result) > 0, f"Filter {f} returned empty"

    def test_rotate_image(self):
        import cv2
        from app.services.doc_scanner import rotate_image
        data = _jpg(400, 300)
        rotated = rotate_image(data, 90)
        img = cv2.imdecode(np.frombuffer(rotated, np.uint8), cv2.IMREAD_COLOR)
        assert img.shape[1] == 300  # width becomes height
        assert img.shape[0] == 400

    def test_scan_document_with_corners(self):
        from app.services.doc_scanner import scan_document
        data = _jpg(800, 600)
        result = scan_document(data, corners=[[50,50],[750,50],[750,550],[50,550]])
        assert result["image"] is not None
        assert result["filter_applied"] == "auto"

    def test_scan_document_no_corners(self):
        from app.services.doc_scanner import scan_document
        data = _jpg(400, 300)
        result = scan_document(data, auto_detect=False)
        assert result["image"] is not None


# ══════════════════════════════════════════════════════════════
# 圖片批次處理
# ══════════════════════════════════════════════════════════════

class TestImageBatch:

    def test_resize_fit(self):
        from app.services.image_batch import resize_image
        result = resize_image(_jpg(800, 600), 200, 200, mode="fit")
        from PIL import Image
        img = Image.open(io.BytesIO(result))
        assert img.size == (200, 200)

    def test_resize_cover(self):
        from app.services.image_batch import resize_image
        result = resize_image(_jpg(800, 600), 200, 200, mode="cover")
        from PIL import Image
        img = Image.open(io.BytesIO(result))
        assert img.size == (200, 200)

    def test_convert_to_png(self):
        from app.services.image_batch import convert_format
        result = convert_format(_jpg(), "PNG")
        assert result[:4] == b'\x89PNG'

    def test_compress(self):
        from app.services.image_batch import compress_image
        original = _jpg(800, 600)
        compressed = compress_image(original, quality=30)
        assert len(compressed) < len(original)

    def test_text_watermark(self):
        from app.services.image_batch import add_text_watermark
        result = add_text_watermark(_jpg(), "TEST", font_size=24, opacity=50)
        assert len(result) > 0

    def test_batch_compress(self):
        from app.services.image_batch import batch_compress
        images = [("a.jpg", _jpg()), ("b.jpg", _jpg(200, 150))]
        result = batch_compress("test_task", images, quality=50)
        import zipfile
        zf = zipfile.ZipFile(io.BytesIO(result))
        assert len(zf.namelist()) == 2


# ══════════════════════════════════════════════════════════════
# PDF 處理
# ══════════════════════════════════════════════════════════════

class TestPdfProcessor:

    def test_merge_pdfs(self):
        from app.services.pdf_processor import merge_pdfs
        pdfs = [("a.pdf", _pdf("Page 1")), ("b.pdf", _pdf("Page 2"))]
        result = merge_pdfs("test", pdfs)
        assert result[:4] == b"%PDF"
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(result))
        assert len(reader.pages) == 2

    def test_text_watermark(self):
        from app.services.pdf_processor import add_text_watermark_to_pdf
        result = add_text_watermark_to_pdf(_pdf(), "DRAFT")
        assert result[:4] == b"%PDF"

    def test_protect_pdf(self):
        from app.services.pdf_processor import protect_pdf
        result = protect_pdf(_pdf(), "secret123")
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(result))
        assert reader.is_encrypted

    def test_pdf_info(self):
        from app.services.pdf_processor import get_pdf_info
        info = get_pdf_info(_pdf())
        assert info["pages"] == 1
        assert not info["encrypted"]


# ══════════════════════════════════════════════════════════════
# 文件轉檔
# ══════════════════════════════════════════════════════════════

class TestDocConverter:

    def test_word_to_pdf(self):
        from app.services.doc_converter import word_to_pdf
        result = word_to_pdf(_docx())
        assert result[:4] == b"%PDF"

    def test_pdf_to_word(self):
        from app.services.doc_converter import pdf_to_word
        result = pdf_to_word(_pdf("Hello World"))
        assert len(result) > 0
        # DOCX 的 magic bytes (ZIP)
        assert result[:2] == b"PK"

    def test_markdown_to_pdf(self):
        from app.services.doc_converter import markdown_to_pdf
        result = markdown_to_pdf("# Title\n\nHello **world**")
        assert result[:4] == b"%PDF"

    def test_markdown_to_word(self):
        from app.services.doc_converter import markdown_to_word
        result = markdown_to_word("# Title\n\n- item1\n- item2")
        assert result[:2] == b"PK"

    def test_word_to_markdown(self):
        from app.services.doc_converter import word_to_markdown
        result = word_to_markdown(_docx("My Heading"))
        assert "My Heading" in result


# ══════════════════════════════════════════════════════════════
# GIF 製作
# ══════════════════════════════════════════════════════════════

class TestGifCreator:

    def test_create_gif(self):
        from app.services.gif_creator import create_gif_from_images
        images = [("1.jpg", _jpg(200, 150, (255,0,0))),
                  ("2.jpg", _jpg(200, 150, (0,255,0))),
                  ("3.jpg", _jpg(200, 150, (0,0,255)))]
        result = create_gif_from_images("test", images, duration_ms=200)
        assert result[:6] in (b"GIF89a", b"GIF87a")

    def test_create_gif_with_resize(self):
        from app.services.gif_creator import create_gif_from_images
        images = [("a.jpg", _jpg(800, 600)), ("b.jpg", _jpg(400, 300))]
        result = create_gif_from_images("test", images, resize_width=100, resize_height=75)
        assert len(result) > 0


# ══════════════════════════════════════════════════════════════
# 批次改名
# ══════════════════════════════════════════════════════════════

class TestBatchRenamer:

    def test_preview_prefix_suffix(self):
        from app.services.batch_renamer import preview_rename
        result = preview_rename(["a.jpg", "b.jpg"], prefix="IMG_", suffix="_v2")
        assert result[0]["renamed"] == "IMG_a_v2.jpg"
        assert result[1]["renamed"] == "IMG_b_v2.jpg"

    def test_preview_find_replace(self):
        from app.services.batch_renamer import preview_rename
        result = preview_rename(["photo_001.jpg"], find="photo", replace="img")
        assert result[0]["renamed"] == "img_001.jpg"

    def test_preview_numbering(self):
        from app.services.batch_renamer import preview_rename
        result = preview_rename(["a.jpg", "b.jpg", "c.jpg"],
                                numbering=True, numbering_start=5, numbering_digits=4)
        assert result[0]["renamed"] == "0005_a.jpg"
        assert result[2]["renamed"] == "0007_c.jpg"

    def test_apply_rename(self):
        from app.services.batch_renamer import apply_rename, preview_rename
        files = [("a.jpg", b"data_a"), ("b.jpg", b"data_b")]
        rename_map = preview_rename(["a.jpg", "b.jpg"], prefix="new_")
        result = apply_rename("test", files, rename_map)
        import zipfile
        zf = zipfile.ZipFile(io.BytesIO(result))
        names = zf.namelist()
        assert "new_a.jpg" in names
        assert "new_b.jpg" in names


# ══════════════════════════════════════════════════════════════
# 群組 + 模板 + 多頁 PDF
# ══════════════════════════════════════════════════════════════

class TestModels:

    @pytest.fixture(autouse=True)
    def setup_db(self, tmp_path, monkeypatch):
        db_path = str(tmp_path / "test.db")
        monkeypatch.setattr("app.config.get_settings",
                            lambda: type("S", (), {"DATABASE_PATH": db_path})())
        from app.database import init_db
        init_db()

    def test_group_crud(self):
        from app.models.group import GroupModel
        from app.models.contact import ContactModel
        c1 = ContactModel.create("u1", "Alice", "alice@test.com")
        c2 = ContactModel.create("u1", "Bob", "bob@test.com")
        gid = GroupModel.create("u1", "Team A", "Description")
        GroupModel.set_members(gid, [c1, c2])

        groups = GroupModel.list_by_user("u1")
        assert len(groups) == 1
        assert groups[0]["member_count"] == 2

        members = GroupModel.get_members(gid)
        assert len(members) == 2
        names = {m["name"] for m in members}
        assert names == {"Alice", "Bob"}

        GroupModel.delete(gid)
        assert GroupModel.list_by_user("u1") == []

    def test_template_defaults(self):
        from app.models.template import TemplateModel
        templates = TemplateModel.list_by_user("u1")
        assert len(templates) == 8  # 8 種文件類型預設模板
        doc_types = {t["doc_type"] for t in templates}
        assert "exam" in doc_types
        assert "official" in doc_types

    def test_template_apply(self):
        from app.models.template import TemplateModel
        t = TemplateModel.get_default("exam")
        result = TemplateModel.apply_template(t, {
            "extracted_text_summary": "材料力學期中考",
            "doc_type_label": "考卷"
        })
        assert "材料力學期中考" in result["subject"]
        assert "材料力學期中考" in result["body"]

    def test_template_custom_override(self):
        from app.models.template import TemplateModel
        TemplateModel.create("u1", "exam", "My Template", "[{doc_type}] {summary}", "Body: {summary}")
        t = TemplateModel.get_for_doc_type("u1", "exam")
        assert t["name"] == "My Template"

    def test_images_to_pdf_multipage(self):
        from app.services.image_processor import images_to_pdf
        pages = [_jpg(400, 300), _jpg(300, 400)]
        result = images_to_pdf(pages)
        assert result[:4] == b"%PDF"
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(result))
        assert len(reader.pages) == 2


# ══════════════════════════════════════════════════════════════
# 背景任務系統
# ══════════════════════════════════════════════════════════════

class TestTaskSystem:

    def test_submit_and_complete(self):
        import time
        from app.core.tasks import submit_task, get_task, update_task_progress

        def my_task(task_id, data):
            update_task_progress(task_id, 50, "working")
            return data * 2

        tid = submit_task(my_task, 21)
        time.sleep(1)
        task = get_task(tid)
        assert task is not None
        assert task.status.value == "completed"
        assert task.result == 42

    def test_submit_failing_task(self):
        import time
        from app.core.tasks import submit_task, get_task

        def bad_task(task_id):
            raise ValueError("boom")

        tid = submit_task(bad_task)
        time.sleep(1)
        task = get_task(tid)
        assert task.status.value == "failed"
        assert "boom" in task.error


# ══════════════════════════════════════════════════════════════
# API 端點整合測試
# ══════════════════════════════════════════════════════════════

class TestAPI:

    @pytest.fixture(autouse=True)
    def setup(self, tmp_path, monkeypatch):
        db_path = str(tmp_path / "test.db")
        monkeypatch.setattr("app.config.get_settings",
                            lambda: type("S", (), {"DATABASE_PATH": db_path})())
        from app.database import init_db
        init_db()

    @pytest.fixture
    def client(self):
        from fastapi.testclient import TestClient
        from main import app
        return TestClient(app)

    def test_health(self, client):
        r = client.get("/health")
        assert r.status_code == 200
        assert "version" in r.json()

    def test_upload_and_detect(self, client):
        r = client.post("/api/upload", files={"file": ("t.jpg", _jpg(), "image/jpeg")})
        assert r.status_code == 200
        r = client.post("/api/scan/detect", json={})
        assert r.status_code == 200
        assert "corners" in r.json()

    def test_upload_and_process(self, client):
        client.post("/api/upload", files={"file": ("t.jpg", _jpg(800,600), "image/jpeg")})
        r = client.post("/api/scan/process", json={
            "corners": [[50,50],[750,50],[750,550],[50,550]],
            "filter_name": "auto", "auto_detect": False
        })
        assert r.status_code == 200
        assert r.json()["success"]

    def test_image_tools_compress(self, client):
        r = client.post("/api/tools/image/compress",
            files={"file": ("t.jpg", _jpg(), "image/jpeg")},
            data={"quality": 50})
        assert r.status_code == 200

    def test_pdf_tools_info(self, client):
        r = client.post("/api/tools/pdf/info",
            files={"file": ("t.pdf", _pdf(), "application/pdf")})
        assert r.status_code == 200
        assert r.json()["pages"] == 1

    def test_doc_convert_md_to_pdf(self, client):
        r = client.post("/api/tools/convert/md-to-pdf",
            files={"file": ("t.md", b"# Hello\n\nWorld", "text/markdown")})
        assert r.status_code == 200

    def test_groups_crud(self, client):
        client.post("/api/contacts", json={"name":"A","email":"a@t.com"})
        contacts = client.get("/api/contacts").json()

        r = client.post("/api/groups", json={"name":"G1","contact_ids":[contacts[0]["id"]]})
        assert r.status_code == 200
        gid = r.json()["id"]

        r = client.get(f"/api/groups/{gid}")
        assert len(r.json()["members"]) == 1

        r = client.delete(f"/api/groups/{gid}")
        assert r.status_code == 200

    def test_templates(self, client):
        r = client.get("/api/templates")
        assert r.status_code == 200
        assert len(r.json()) >= 8

    def test_rename_preview(self, client):
        r = client.post("/api/tools/rename/preview", json={
            "filenames": ["a.jpg", "b.jpg"],
            "prefix": "new_"
        })
        assert r.status_code == 200
        assert r.json()["results"][0]["renamed"] == "new_a.jpg"
