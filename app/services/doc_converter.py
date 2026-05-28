"""文件格式轉換 — Word⟷PDF、Markdown⟷PDF/Word

純 Python 實作，不依賴 LibreOffice。
"""
import glob
import io
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)


# ══════════════════════════════════════════════════════════════
# 1. Word → PDF
# ══════════════════════════════════════════════════════════════

def word_to_pdf(docx_data: bytes) -> bytes:
    """將 DOCX 轉為 PDF

    使用 pymupdf 的 Document 支援 + python-docx 提取內容後
    透過 ReportLab 渲染為 PDF。
    """
    from docx import Document
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 嘗試註冊中文字型
    _register_cjk_font()

    doc = Document(io.BytesIO(docx_data))
    buf = io.BytesIO()

    pdf = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=25*mm, rightMargin=25*mm,
                            topMargin=20*mm, bottomMargin=20*mm)

    styles = getSampleStyleSheet()
    # 建立支援中文的樣式
    body_font = _get_available_font()
    body_style = ParagraphStyle(
        "CJKBody", parent=styles["Normal"],
        fontName=body_font, fontSize=11, leading=18,
    )
    heading_style = ParagraphStyle(
        "CJKHeading", parent=styles["Heading1"],
        fontName=body_font, fontSize=16, leading=24,
        spaceAfter=12,
    )

    story = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 6))
            continue

        # 簡單判斷是否為標題
        if para.style and para.style.name and "Heading" in para.style.name:
            story.append(Paragraph(_escape_xml(text), heading_style))
        else:
            story.append(Paragraph(_escape_xml(text), body_style))

    if not story:
        story.append(Paragraph("(空白文件)", body_style))

    pdf.build(story)
    result = buf.getvalue()
    logger.info("Word → PDF 完成: %d bytes", len(result))
    return result


# ══════════════════════════════════════════════════════════════
# 2. PDF → Word
# ══════════════════════════════════════════════════════════════

def pdf_to_word(pdf_data: bytes) -> bytes:
    """將 PDF 轉為 DOCX（提取文字內容）"""
    import fitz  # pymupdf
    from docx import Document
    from docx.shared import Pt

    pdf_doc = fitz.open(stream=pdf_data, filetype="pdf")
    word_doc = Document()

    for i, page in enumerate(pdf_doc):
        if i > 0:
            word_doc.add_page_break()

        text = page.get_text("text")
        for line in text.split("\n"):
            line = line.strip()
            if line:
                p = word_doc.add_paragraph(line)
                p.style.font.size = Pt(11)

    buf = io.BytesIO()
    word_doc.save(buf)
    result = buf.getvalue()
    logger.info("PDF → Word 完成: %d 頁, %d bytes", len(pdf_doc), len(result))
    return result


# ══════════════════════════════════════════════════════════════
# 3. Markdown → PDF
# ══════════════════════════════════════════════════════════════

def markdown_to_pdf(md_text: str) -> bytes:
    """將 Markdown 轉為 PDF"""
    import markdown
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    _register_cjk_font()
    body_font = _get_available_font()

    # Markdown → HTML
    html = markdown.markdown(md_text, extensions=["tables", "fenced_code"])

    # HTML → 簡單文字段落（用 BeautifulSoup 提取）
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=25*mm, rightMargin=25*mm,
                            topMargin=20*mm, bottomMargin=20*mm)

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle("MDBody", parent=styles["Normal"],
                                fontName=body_font, fontSize=11, leading=18)
    h1_style = ParagraphStyle("MDH1", parent=styles["Heading1"],
                              fontName=body_font, fontSize=20, leading=28, spaceAfter=12)
    h2_style = ParagraphStyle("MDH2", parent=styles["Heading2"],
                              fontName=body_font, fontSize=16, leading=22, spaceAfter=10)
    code_style = ParagraphStyle("MDCode", parent=styles["Code"],
                                fontName="Courier", fontSize=9, leading=12,
                                leftIndent=20, spaceAfter=8)

    story = []
    for element in soup.children:
        tag = getattr(element, "name", None)
        text = element.get_text(strip=True) if hasattr(element, "get_text") else str(element).strip()
        if not text:
            continue

        if tag == "h1":
            story.append(Paragraph(_escape_xml(text), h1_style))
        elif tag == "h2":
            story.append(Paragraph(_escape_xml(text), h2_style))
        elif tag in ("h3", "h4", "h5", "h6"):
            story.append(Paragraph(_escape_xml(text), h2_style))
        elif tag in ("pre", "code"):
            story.append(Paragraph(_escape_xml(text), code_style))
        elif tag:
            story.append(Paragraph(_escape_xml(text), body_style))
        else:
            if text:
                story.append(Paragraph(_escape_xml(text), body_style))

    if not story:
        story.append(Paragraph("(空白文件)", body_style))

    pdf.build(story)
    result = buf.getvalue()
    logger.info("Markdown → PDF 完成: %d bytes", len(result))
    return result


# ══════════════════════════════════════════════════════════════
# 4. Markdown → Word
# ══════════════════════════════════════════════════════════════

def markdown_to_word(md_text: str) -> bytes:
    """將 Markdown 轉為 DOCX"""
    import markdown
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Pt

    html = markdown.markdown(md_text, extensions=["tables", "fenced_code"])
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()

    for element in soup.children:
        tag = getattr(element, "name", None)
        text = element.get_text(strip=True) if hasattr(element, "get_text") else str(element).strip()
        if not text:
            continue

        if tag == "h1":
            doc.add_heading(text, level=1)
        elif tag == "h2":
            doc.add_heading(text, level=2)
        elif tag in ("h3", "h4", "h5", "h6"):
            doc.add_heading(text, level=3)
        elif tag in ("ul", "ol"):
            for li in element.find_all("li"):
                li_text = li.get_text(strip=True)
                if li_text:
                    doc.add_paragraph(li_text, style="List Bullet")
        elif tag in ("pre", "code"):
            p = doc.add_paragraph(text)
            p.style.font.size = Pt(9)
            p.style.font.name = "Courier New"
        else:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    result = buf.getvalue()
    logger.info("Markdown → Word 完成: %d bytes", len(result))
    return result


# ══════════════════════════════════════════════════════════════
# 5. Word → Markdown
# ══════════════════════════════════════════════════════════════

def word_to_markdown(docx_data: bytes) -> str:
    """將 DOCX 轉為 Markdown 文字"""
    from docx import Document

    doc = Document(io.BytesIO(docx_data))
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        style_name = para.style.name if para.style else ""

        if "Heading 1" in style_name:
            lines.append(f"# {text}")
        elif "Heading 2" in style_name:
            lines.append(f"## {text}")
        elif "Heading 3" in style_name:
            lines.append(f"### {text}")
        elif "List" in style_name:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    result = "\n\n".join(lines)
    logger.info("Word → Markdown 完成: %d chars", len(result))
    return result


# ══════════════════════════════════════════════════════════════
# 6. PDF → Markdown
# ══════════════════════════════════════════════════════════════

def pdf_to_markdown(pdf_data: bytes) -> str:
    """將 PDF 轉為 Markdown 文字

    策略：
    1. 用 pymupdf 抽出每行文字 + 平均字體大小
    2. 統計全文字體大小分布，把最大的 1~2 個 size 視為 H1 / H2
    3. 其餘為內文段落
    4. 連續空行壓成一個段落分隔
    """
    import fitz  # pymupdf

    pdf_doc = fitz.open(stream=pdf_data, filetype="pdf")

    # 第一遍：收集所有 line + avg font size
    all_lines: list[tuple[str, float, bool]] = []  # (text, size, is_bold)
    size_counts: dict[int, int] = {}

    for page in pdf_doc:
        blocks = page.get_text("dict").get("blocks", [])
        for block in blocks:
            if block.get("type") != 0:  # 0 = text block
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                # 行內所有 span 的字級平均（通常一行字級一致）
                sizes = [s.get("size", 11) for s in spans]
                avg_size = sum(sizes) / len(sizes)
                # 任一 span 是粗體就算粗體
                is_bold = any("Bold" in (s.get("font", "") or "") for s in spans)
                text = "".join(s.get("text", "") for s in spans).strip()
                if not text:
                    continue
                all_lines.append((text, avg_size, is_bold))
                bucket = round(avg_size)
                size_counts[bucket] = size_counts.get(bucket, 0) + 1
        all_lines.append(("", 0.0, False))  # 頁尾標記分段

    pdf_doc.close()

    # 決定 heading 閾值：取最常見字級為 body，超過 body 的視為 heading
    if size_counts:
        body_size = max(size_counts.items(), key=lambda x: x[1])[0]
    else:
        body_size = 11
    h1_threshold = body_size + 6
    h2_threshold = body_size + 3
    h3_threshold = body_size + 1

    # 第二遍：組 markdown
    lines: list[str] = []
    prev_blank = True
    for text, size, is_bold in all_lines:
        if not text:
            if not prev_blank:
                lines.append("")
                prev_blank = True
            continue
        prev_blank = False
        if size >= h1_threshold:
            lines.append(f"# {text}")
        elif size >= h2_threshold:
            lines.append(f"## {text}")
        elif size >= h3_threshold or is_bold:
            lines.append(f"### {text}" if size >= h3_threshold else f"**{text}**")
        else:
            lines.append(text)

    # 收尾：去掉開頭/結尾多餘空行
    result = "\n".join(lines).strip() + "\n"
    logger.info("PDF → Markdown 完成: %d chars (body_size=%d)", len(result), body_size)
    return result


# ══════════════════════════════════════════════════════════════
# 共用工具
# ══════════════════════════════════════════════════════════════

def _escape_xml(text: str) -> str:
    """轉義 XML 特殊字元（ReportLab Paragraph 需要）"""
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;"))


_cjk_registered = False

def _cjk_font_candidates() -> list[str]:
    """跨平台 CJK 字型候選清單（依優先序）。

    刻意**不**納入 DejaVu 等無中文 glyph 的字型——那會「註冊成功卻把中文
    渲染成空白（\\x00）」，反而遮蔽問題（即 CI 上中文測試失敗的真因）。
    """
    paths = [
        # Linux：對應 fonts-noto-cjk / fonts-wqy-zenhei / fonts-arphic-uming
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/truetype/arphic/uming.ttc",
    ]
    # Windows：新細明 / 細明 / 微軟正黑 / 標楷（皆可被 ReportLab 內嵌並含中文）
    windir = os.environ.get("WINDIR", r"C:\Windows")
    paths += [os.path.join(windir, "Fonts", f)
              for f in ("msjh.ttc", "mingliu.ttc", "simsun.ttc", "kaiu.ttf")]
    # macOS
    paths += [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc",
    ]
    # 萬用 glob：吃進不同套件版本造成的檔名差異（如 NotoSansCJK-VF.otf.ttc）
    for pat in (
        "/usr/share/fonts/**/NotoSansCJK*.tt[cf]",
        "/usr/share/fonts/**/NotoSerifCJK*.tt[cf]",
        "/usr/share/fonts/**/wqy-*.ttc",
    ):
        paths += sorted(glob.glob(pat, recursive=True))
    return paths


def _register_cjk_font():
    """註冊系統中第一個可用的 CJK 字型（跨 Linux/Windows/macOS）"""
    global _cjk_registered
    if _cjk_registered:
        return

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    seen = set()
    for path in _cjk_font_candidates():
        if path in seen or not os.path.exists(path):
            continue
        seen.add(path)
        try:
            pdfmetrics.registerFont(TTFont("CJKFont", path))
            _cjk_registered = True
            logger.info("已註冊 CJK 字型: %s", path)
            return
        except Exception:
            continue  # 非 TrueType / 無法內嵌 → 換下一個候選

    logger.warning("找不到可用的 CJK 字型，中文可能無法正確渲染")
    _cjk_registered = True  # 標記為已嘗試，避免重複


def _get_available_font() -> str:
    """取得可用的字型名稱"""
    from reportlab.pdfbase import pdfmetrics
    if "CJKFont" in pdfmetrics.getRegisteredFontNames():
        return "CJKFont"
    return "Helvetica"


# ─── 公開介面 ───────────────────────────────────────────────
# 給其他模組（例如 form_fill.filler）使用，避免直接 import 私有名稱

def ensure_cjk_font() -> str:
    """確保已註冊 CJK 字型並回傳可用的 ReportLab 字型名稱。

    若系統有 NotoSansCJK / WenQuanYi / 新細明等含中文 glyph 的字型則回
    "CJKFont"，否則 fallback 為 "Helvetica"（無中文支援）。
    """
    _register_cjk_font()
    return _get_available_font()
